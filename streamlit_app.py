# streamlit_app.py
import streamlit as st
import tempfile
import os
import importlib
import inspect
import zipfile
import io
import threading
import time
import random
from typing import List, Callable, Any

# Import Document to read docx text for Diff View
from docx import Document

from llm_client import LLMClient
from doc_processor import DocProcessor
from dotenv import load_dotenv

# Load environment variables from .env if present
load_dotenv()

# Config: module name for MCP tools (default: mcp_rules.py)
MCP_MODULE_NAME = os.environ.get("MCP_MODULE", "mcp_rules")

# LLM model to use for grammar correction
OPENAI_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4o-mini")

st.set_page_config(page_title="DOCX: MCP-only Rules + Grammar Processor", layout="wide")
st.title("DOCX Processor — MCP-only custom rules + LLM grammar")
st.write("This app applies **only** MCP-provided transformation tools (from `mcp_rules.py`) and optional LLM grammar correction. "
         "Local replacement rules are intentionally excluded.")


# --- Helper for Diff View ---
def extract_text_from_docx(path: str) -> str:
    """Extracts text from paragraphs and tables for comparison."""
    try:
        doc = Document(path)
        full_text = []
        
        # Extract body paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract table text (simple linear extraction)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text)
                            
        return "\n\n".join(full_text)
    except Exception as e:
        return f"[Error extraction text: {str(e)}]"


# Discover MCP tool functions from the MCP module (functions that end with '_tool')
mcp_funcs: List[Callable[[str], Any]] = []
mcp_available = False
mcp_import_errors = None
mcp_mod = None

try:
    # reload if already imported (development friendliness)
    if MCP_MODULE_NAME in importlib.sys.modules:
        importlib.reload(importlib.import_module(MCP_MODULE_NAME))
    mcp_mod = importlib.import_module(MCP_MODULE_NAME)

    # collect async and sync functions named *_tool
    for name, obj in inspect.getmembers(mcp_mod, inspect.iscoroutinefunction):
        if name.endswith("_tool"):
            mcp_funcs.append(obj)
    for name, obj in inspect.getmembers(mcp_mod, inspect.isfunction):
        if name.endswith("_tool") and obj not in mcp_funcs:
            mcp_funcs.append(obj)

    if mcp_funcs:
        mcp_available = True
except Exception as e:
    mcp_import_errors = str(e)
    mcp_mod = None

# UI: show discovered tools and options on left, file upload on right

with st.sidebar:
    st.header("MCP Rules (tools)")
    if mcp_available:
        st.success(f"Found MCP module '{MCP_MODULE_NAME}' with {len(mcp_funcs)} tool(s).")
        st.write("Discovered MCP tool functions (executed sequentially on each text block):")
        # show list inside a scrollable box (fixed max-height)
        rules_html = "<div style='max-height:300px; overflow:auto; padding:6px 8px; border:1px solid #eee; border-radius:4px;'>"
        for f in mcp_funcs:
            rules_html += f"<div style='margin:4px 0;'>• <strong>{f.__name__}</strong></div>"
        rules_html += "</div>"
        st.markdown(rules_html, unsafe_allow_html=True)
        st.markdown("---")
        st.info("Only MCP transforms from the above functions will be applied. Local replacement rules are disabled in this UI.")
    else:
        st.warning("No MCP module loaded. Place 'mcp_rules.py' in the app folder (or set MCP_MODULE env var).")
        if mcp_import_errors:
            st.write("Import error:", mcp_import_errors)
        st.write("You can still use the LLM grammar correction step if configured.")
    st.markdown("---")
    st.header("Options")
    apply_mcp = st.checkbox("Apply MCP transforms", value=True, help="If disabled, MCP tools will not be run.")
    apply_grammar = st.checkbox("Apply LLM grammar correction (requires OPENAI_API_KEY)", value=False)
    apply_mcp_first = st.checkbox("Apply MCP transforms before LLM grammar correction", value=True)


cols = st.columns([1, 2, 1])
with cols[1]:
    st.header("Process .docx Files")
    uploaded_files = st.file_uploader("Upload .docx", type=["docx"], accept_multiple_files=True)
    
    if uploaded_files:
        count = len(uploaded_files)
        st.info(f"Uploaded {count} file{'s' if count > 1 else ''}")
        
        # Prepare file metadata for processing
        input_files_map = [] 
        
        for uploaded_file in uploaded_files:
            # Save upload to temp file
            t_in = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            t_in.write(uploaded_file.read())
            t_in.flush()
            t_in.close()
            
            # Prepare output temp file
            t_out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            t_out_path = t_out.name
            t_out.close()
            
            input_files_map.append({
                "input_path": t_in.name,
                "output_path": t_out_path,
                "original_name": uploaded_file.name
            })

        if st.button(f"Process {count} file{'s' if count > 1 else ''}"):

            # Prepare transform functions
            transforms = []
            if apply_mcp and mcp_available:
                transforms = list(mcp_funcs)

            # Prepare LLM client
            llm = None
            if apply_grammar:
                if os.getenv("OPENAI_API_KEY"):
                    llm = LLMClient(model=OPENAI_MODEL)
                else:
                    st.warning("OPENAI_API_KEY not set — grammar step will be skipped.")
                    llm = None

            dp = DocProcessor(
                rules_manager=None, 
                llm_client=llm,
                apply_grammar=apply_grammar,
                mcp_transform_funcs=transforms
            )

            # Prepare progress UI
            progress_bar = st.progress(0)
            status = st.empty()
            status.info("Starting processing...")
            result = {"error": None}

            def _run_processing():
                try:
                    for file_info in input_files_map:
                        dp.process_docx(
                            file_info["input_path"], 
                            file_info["output_path"], 
                            apply_rules_first=apply_mcp_first
                        )
                except Exception as e:
                    result["error"] = e

            # Run processing in a background thread
            worker = threading.Thread(target=_run_processing, daemon=True)
            worker.start()

            # Animate progress
            pct = 0
            last_update = time.time()
            while worker.is_alive():
                if time.time() - last_update > 0.25:
                    increment = random.randint(2, 6)
                    pct = min(pct + increment, 95)
                    progress_bar.progress(pct)
                    last_update = time.time()
                time.sleep(0.1)

            worker.join()

            if result["error"] is None:
                progress_bar.progress(100)
                status.success("Processing finished.")
                
                # --- Diff View Section ---
                st.subheader("Comparison & Review")
                for f_info in input_files_map:
                    with st.expander(f"Review: {f_info['original_name']}", expanded=False):
                        orig_text = extract_text_from_docx(f_info["input_path"])
                        proc_text = extract_text_from_docx(f_info["output_path"])
                        
                        diff_cols = st.columns(2)
                        with diff_cols[0]:
                            st.markdown("**Original Text**")
                            st.text_area(label="orig", value=orig_text, height=300, disabled=True, key=f"orig_{f_info['input_path']}")
                        with diff_cols[1]:
                            st.markdown("**Processed Text**")
                            st.text_area(label="proc", value=proc_text, height=300, disabled=True, key=f"proc_{f_info['output_path']}")

                # --- Download Section ---
                st.subheader("Download Results")
                try:
                    # Single file download
                    if len(input_files_map) == 1:
                        f_info = input_files_map[0]
                        with open(f_info["output_path"], "rb") as f:
                            data = f.read()
                            st.download_button(
                                "Download processed .docx",
                                data=data,
                                file_name=f"processed_{f_info['original_name']}",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                    
                    # Batch ZIP download
                    else:
                        zip_buffer = io.BytesIO()
                        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                            for f_info in input_files_map:
                                zf.write(f_info["output_path"], f"processed_{f_info['original_name']}")
                        
                        st.download_button(
                            "Download all files (ZIP)",
                            data=zip_buffer.getvalue(),
                            file_name="processed_files.zip",
                            mime="application/zip"
                        )

                except Exception as e:
                    st.error(f"Failed to prepare download: {e}")
            else:
                progress_bar.progress(100)
                st.error(f"Processing failed: {result['error']}")

            # Cleanup
            for f_info in input_files_map:
                try:
                    if os.path.exists(f_info["input_path"]):
                        os.unlink(f_info["input_path"])
                    if os.path.exists(f_info["output_path"]):
                        os.unlink(f_info["output_path"])
                except Exception:
                    pass

    else:
        st.info("Upload .docx file(s) to enable processing options.")

st.markdown("---")
st.write("Notes:")
st.write("- **Diff View** extracts text from document body and tables for a quick side-by-side comparison.")
st.write("- This app intentionally excludes local replacement rules; only MCP tool functions (from `mcp_rules.py`) are applied.")
st.write("- LLM grammar correction is applied per text block and may incur API costs.")