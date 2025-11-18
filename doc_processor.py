# doc_processor.py
from docx import Document
from typing import Optional, Callable, List, Any
import asyncio

from rules_manager import RulesManager
from llm_client import LLMClient

class DocProcessor:
    def __init__(
        self,
        rules_manager: Optional[RulesManager] = None,
        llm_client: Optional[LLMClient] = None,
        apply_grammar: bool = True,
        mcp_transform_funcs: Optional[List[Callable[[str], Any]]] = None,
    ):
        """
        mcp_transform_funcs: list of callables (sync or async) accepting text and returning transformed text.
        """
        self.rules_manager = rules_manager
        self.llm_client = llm_client
        self.apply_grammar = apply_grammar
        self.mcp_transform_funcs = mcp_transform_funcs or []

    def _apply_transforms(self, text: str) -> str:
        out = text
        for func in self.mcp_transform_funcs:
            try:
                if asyncio.iscoroutinefunction(func):
                    # run async function in a new event loop
                    out = asyncio.run(func(out))
                else:
                    out = func(out)
            except Exception as e:
                print(f"Transform function {getattr(func,'__name__',str(func))} failed: {e}")
                continue
            if out is None:
                out = ""
            else:
                out = str(out)
        return out

    def _apply_local_rules(self, text: str) -> str:
        if self.rules_manager:
            return self.rules_manager.apply_rules_to_text(text)
        return text

    def _apply_llm(self, text: str) -> str:
        if self.apply_grammar and self.llm_client:
            try:
                return self.llm_client.correct_grammar(text)
            except Exception as e:
                print("LLM correction failed:", e)
                return text
        return text

    def process_text(self, orig_text: str, apply_rules_first: bool = True) -> str:
        text = orig_text
        if apply_rules_first:
            text = self._apply_transforms(text)
            text = self._apply_local_rules(text)
            text = self._apply_llm(text)
        else:
            text = self._apply_llm(text)
            text = self._apply_transforms(text)
            text = self._apply_local_rules(text)
        return text

    def process_docx(self, input_path: str, output_path: str, apply_rules_first: bool = True):
        doc = Document(input_path)

        for paragraph in doc.paragraphs:
            original = paragraph.text
            processed = self.process_text(original, apply_rules_first)
            if processed != original:
                self._replace_paragraph_text(paragraph, processed)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        original = paragraph.text
                        processed = self.process_text(original, apply_rules_first)
                        if processed != original:
                            self._replace_paragraph_text(paragraph, processed)

        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                original = paragraph.text
                processed = self.process_text(original, apply_rules_first)
                if processed != original:
                    self._replace_paragraph_text(paragraph, processed)
            footer = section.footer
            for paragraph in footer.paragraphs:
                original = paragraph.text
                processed = self.process_text(original, apply_rules_first)
                if processed != original:
                    self._replace_paragraph_text(paragraph, processed)

        doc.save(output_path)

    def _replace_paragraph_text(self, paragraph, new_text: str):
        # Simple run-replace (may change run-level formatting)
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)
